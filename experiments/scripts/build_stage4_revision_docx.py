"""Build a Word report for the 24h MMdedup Plan B experiment closure."""

from __future__ import annotations

import csv
import json
import statistics
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
RESULT_ROOT = ROOT / "experiments/results/plan_b_stage4"
REPORT_DIR = ROOT / "docs/reports"
DOCX_PATH = REPORT_DIR / "mmdedup_pipeline_revision_experiment_plan.docx"


def main() -> int:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    data = load_data()
    doc = Document()
    setup_document(doc)
    add_cover(doc, data)
    add_pipeline_changes(doc)
    add_stage4_method(doc)
    add_experiments(doc, data)
    add_source_rule(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)
    return 0


def load_data() -> dict[str, Any]:
    return {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "audit": read_json("data_audits/2026-05-20_data_reasonableness_audit.json"),
        "adjudication": read_json("exp_stage4_adjudicated_1000_200k_high_joint_20260519/metrics.json"),
        "stage4_eval": read_json("exp_stage4_eval_1000_200k_high_joint_20260519/metrics.json"),
        "error_analysis": read_json("exp_stage4_error_analysis_1000_200k_high_joint_20260520/metrics.json"),
        "split_metrics": read_json("exp_stage4_training_manifests_200k_20260520/metrics.json"),
        "threshold_rows": read_csv("exp_stage4_eval_1000_200k_high_joint_20260519/per_threshold_metrics.csv"),
        "dedup_rows": read_csv("exp_stage4_split_threshold_200k_20260520/threshold_dedup_rates.csv"),
        "cc3m_prepare": read_json("windows_sync/cc3m_subset_200k_20260515/prepare_metrics.json"),
        "candidate_mining": read_json("windows_sync/exp_stage4_candidates_200k_manifest_20260516/metrics.json"),
        "split_threshold": read_json("exp_stage4_split_threshold_200k_20260520/metrics.json"),
        "llava": load_llava_metrics(),
    }


def read_json(rel_path: str) -> dict[str, Any]:
    with (RESULT_ROOT / rel_path).open(encoding="utf-8") as f:
        return json.load(f)


def read_csv(rel_path: str) -> list[dict[str, str]]:
    with (RESULT_ROOT / rel_path).open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_llava_metrics() -> dict[str, dict[str, Any]]:
    mapping = {
        "A": "windows_sync/exp_llava_stage4_train25k_A_raw_25000_2000steps_20260521/metrics.json",
        "B": "windows_sync/exp_llava_stage4_train25k_B_image_only_25000_2000steps_20260521/metrics.json",
        "C": "windows_sync/exp_llava_stage4_train25k_C_text_only_25000_2000steps_20260521/metrics.json",
        "D": "windows_sync/exp_llava_stage4_train25k_D_naive_union_25000_2000steps_20260521/metrics.json",
    }
    out: dict[str, dict[str, Any]] = {}
    for split, rel_path in mapping.items():
        path = RESULT_ROOT / rel_path
        if path.exists():
            with path.open(encoding="utf-8") as f:
                out[split] = json.load(f)
    return out


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(4)
    for style_name, size, color in [
        ("Title", 22, "172033"),
        ("Heading 1", 16, "172033"),
        ("Heading 2", 13, "1f2937"),
        ("Heading 3", 11.5, "334155"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_cover(doc: Document, data: dict[str, Any]) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MMdedup 去重流水线修改与实验设计报告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(23, 32, 51)
    set_run_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Plan B / Stage 4 image-caption pair-level cross-modal deduplication")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(71, 85, 105)
    set_run_font(run)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"生成时间：{data['generated_at']} Asia/Shanghai")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)
    set_run_font(run)

    add_note_box(
        doc,
        "文档定位",
        [
            "本文件说明本轮对 MMdedup 去重流水线的具体修改，以及 24 小时内需要闭环的实验设计和当前已获得数据。",
            "本文不单独列“待补充数据”章节；E 组训练结果与 VQAv2 下游结果直接在对应实验表格中留空，实验完成后补入。",
        ],
    )


def add_pipeline_changes(doc: Document) -> None:
    add_heading(doc, "1. 本轮对 MMdedup 去重流水线的修改", 1)
    add_paragraph(
        doc,
        "原始 MMdedup 流水线主要由图像、文本、音频三个单模态去重模块组成。"
        "本轮修改保留原有单模态能力，但在图文训练语料场景下新增 Stage 4，"
        "将 image-caption pair 作为 MLLM 训练样本单元进行跨模态去重。"
    )
    add_paragraph(
        doc,
        "修改后的流水线不再只回答“图像是否重复”或“文本是否重复”，而是进一步回答："
        "两个图文对作为训练样本是否构成重复或近重复。这个改动让 MMdedup 的多模态部分从"
        "单模态结果拼接，转向图文对级别的联合判断。"
    )
    add_table(
        doc,
        ["阶段", "原有功能", "本轮修改后的角色"],
        [
            ["Stage 1: Image dedup", "识别重复或近重复图像", "保留为单模态 baseline 和前置清洗信号"],
            ["Stage 2: Audio dedup", "音频重复检测", "保留系统能力，本轮 CIKM 主线不新增音频实验"],
            ["Stage 3: Text dedup", "文本重复检测", "保留为单模态 baseline 和 naive union 组成部分"],
            ["Stage 4: Pair-level dedup", "原流水线中不存在", "新增图文对级别跨模态去重，作为本轮核心修改"],
            ["Source-of-truth", "原稿数字容易不一致", "每个实验记录 ledger、metrics、config、logs，论文数字必须可追溯"],
        ],
    )


def add_stage4_method(doc: Document) -> None:
    add_heading(doc, "2. Stage 4 图文对级别去重方法", 1)
    add_paragraph(
        doc,
        "原稿的 MMdedup 采用 classification-and-clean 思路：先判断文件模态，再分别进入图像、音频、文本去重模块。"
        "这个设计能够解释系统的端到端工程流程，但在 MLLM 图文训练语料场景下仍然存在一个核心缺口："
        "训练样本不是孤立图片或孤立文本，而是 image-caption pair。因此本轮修改新增 Stage 4，"
        "把图文对本身作为去重对象，直接判断两个训练样本单元是否构成重复或近重复。"
    )
    add_heading(doc, "2.1 输入定义与训练样本单元", 2)
    add_paragraph(
        doc,
        "Stage 4 的输入是一组图文对 X = {x_i}，其中每个样本 x_i = (I_i, T_i, id_i)，"
        "I_i 表示图片路径或图片内容，T_i 表示对应 caption，id_i 是可追溯的 pair id。"
        "所有 keep/drop、duplicate group、训练 manifest 和论文统计都以 id_i 为主键。"
        "这样做的好处是：前面图像或文本模块可以继续作为单模态清洗信号，但最终是否删除一个训练样本，"
        "由图文对级别的证据决定，而不是由两个单模态删除集合简单拼接决定。"
    )
    add_table(
        doc,
        ["符号", "含义", "在 Stage 4 中的作用"],
        [
            ["I_i", "第 i 个样本的图片", "输入 CLIP image encoder"],
            ["T_i", "第 i 个样本的 caption", "输入 CLIP text encoder"],
            ["e_i^img", "归一化后的图像 embedding", "计算 image-only 相似度和 joint 表示"],
            ["e_i^txt", "归一化后的文本 embedding", "计算 text-only 相似度和 joint 表示"],
            ["z_i", "图文对 joint embedding", "Stage 4 判重的主表示"],
            ["s_joint(i,j)", "两个图文对的 joint similarity", "与 tau_cross 比较得到重复边"],
            ["a_i", "图片与 caption 的 CLIP alignment", "duplicate group 内选择 keeper"],
        ],
    )
    add_heading(doc, "2.2 图文联合表示", 2)
    add_paragraph(
        doc,
        "对每个图文对，Stage 4 分别使用 CLIP/OpenCLIP 的 image encoder 和 text encoder 获得两路表示。"
        "图像表示和文本表示都会进行 L2 normalization，使后续相似度可以用 cosine similarity 或等价的内积计算。"
        "当前主方案采用 concat joint embedding：z_i = L2Norm([e_i^img; e_i^txt])。"
        "这种表示保留了视觉相似性和文本语义相似性的独立信息，比直接加权求和更容易诊断："
        "当一个错误来自 caption 模板重复、图片实际不同，或者图片近似但 caption 指向不同语义时，"
        "我们仍然可以分别查看 image/text/joint 三类分数。weighted sum 被保留为后续消融方向，"
        "但当前 24 小时闭环中不把它作为主报告数字。"
    )
    add_heading(doc, "2.3 候选挖掘与相似度计算", 2)
    add_paragraph(
        doc,
        "真实 CC3M 中自然重复比例较低，随机抽取 pair-pairs 会得到大量负例，无法有效评估去重算法。"
        "因此 Stage 4 先做 candidate mining：在 200K CC3M 图文对上利用 image similarity、text similarity "
        "和 joint similarity 挖掘高相似候选边，再从这些 hard candidates 中构造人工标注集和评价集。"
        "对候选边 (i,j)，系统记录 s_img(i,j)、s_txt(i,j)、s_joint(i,j) 三类分数。"
        "其中 image-only 和 text-only 用作 baseline，naive union 用作单模态拼接 baseline，"
        "Stage 4 则使用 s_joint(i,j) 作为图文对级别判重的主分数。"
    )
    add_heading(doc, "2.4 判重规则与 keeper 选择", 2)
    add_paragraph(
        doc,
        "Stage 4 将候选边转化为重复图：当 s_joint(i,j) >= tau_cross 时，认为两个图文对之间存在重复或近重复边。"
        "随后在重复图上形成 connected components。每个 component 内只保留一个代表样本，其余样本进入 drops。"
        "keeper 的选择遵循质量优先原则：首先保留 CLIP image-text alignment 更高的图文对，"
        "因为它更可能是图像与 caption 匹配更好的训练样本；如果 alignment 分数非常接近，"
        "再使用图像分辨率、文件大小或稳定 id 顺序作为 tie-breaker。最终输出 keepers、drops、duplicate_groups、"
        "summary.json 和 embedding/cache 路径，保证后续训练 split 与论文数字都可以追溯。"
    )
    add_table(
        doc,
        ["模块", "当前实现", "论文中需要强调的点"],
        [
            ["Pair encoder", "OpenCLIP/CLIP image encoder + text encoder", "图像和文本分别编码，避免把多模态问题退化成单一路径"],
            ["Joint representation", "concat 后 L2 normalization", "保留视觉与文本两路证据，当前主方案可解释性更强"],
            ["Candidate mining", "image/text/joint top-k 候选边", "解决真实数据自然重复太少的问题"],
            ["Duplicate decision", "s_joint >= tau_cross", "判重对象是 image-caption pair，而不是单张图片或单条 caption"],
            ["Keeper selection", "alignment 优先，图像质量和稳定顺序兜底", "删除重复样本时尽量保留图文匹配质量更高的训练样本"],
            ["Reproducible outputs", "keepers/drops/groups/summary/cache", "所有训练 split 和论文表格都能回溯到 source-of-truth"],
        ],
    )
    add_heading(doc, "2.5 与单模态拼接 baseline 的区别", 2)
    add_paragraph(
        doc,
        "naive multimodal union 的逻辑是：image-only 认为重复或 text-only 认为重复，就删除对应样本。"
        "这种方法本质上仍然是两个单模态判断的后处理，容易误删共享 caption 模板但图片不同的样本，"
        "也可能漏掉单模态信号不够强、但图文组合高度相似的样本。Stage 4 的不同之处在于，"
        "它直接学习和比较图文对联合表示，把“是否构成重复训练样本”作为目标。当前 1,000 条人工标注集上，"
        "Stage 4 joint 相比 naive union 有更高 F1，说明它确实提供了比简单拼接更贴近图文训练单元的信号。"
        "同时我们也保留 image-only 对照，因为当前数据中视觉近重复仍然是一个很强的 baseline；"
        "论文写作时应如实呈现这一点，避免把 Stage 4 写成压过所有单模态方法的绝对最优算法。"
    )
    add_heading(doc, "2.6 复杂度与可扩展性", 2)
    add_paragraph(
        doc,
        "如果直接比较所有 pair-pairs，复杂度是 O(n^2)，在 CC3M 规模上不可接受。"
        "因此当前实现沿用原稿中“先向量化、再近邻/聚类收缩候选、最后精细判断”的系统思想："
        "CLIP 编码阶段复杂度约为 O(n*T_CLIP)，候选挖掘阶段只保留 top-k 或超过最小相似度的候选边，"
        "最终判重只在候选集合 E 上执行，复杂度从全量 O(n^2) 降为 O(|E|)。"
        "当前 source-of-truth 中，200K CC3M pool 产生 500K candidate edges，"
        "OpenCLIP candidate mining 在 Windows RTX 3090 上完成，后续 threshold-to-split graph computation "
        "只需要秒级到分钟级开销。因此 Stage 4 的主要成本来自一次性 embedding/candidate mining，"
        "不是下游每个阈值或每个 split 的重复计算。"
    )
    add_heading(doc, "2.7 当前报告采用的关键参数", 2)
    add_table(
        doc,
        ["参数", "当前取值", "说明"],
        [
            ["Data pool", "CC3M 200K image-caption pairs", "用于候选挖掘、split 生成和效率统计"],
            ["Candidate edges", "500K pair-pairs", "由 image/text/joint 相似度挖掘得到"],
            ["Annotation set", "1,000 hard candidates", "人工标注 duplicate / near-duplicate / not-duplicate"],
            ["Positive label", "duplicate + near-duplicate", "用于 P/R/F1 计算"],
            ["Joint embedding", "concat([e_img; e_txt])", "当前 Stage 4 主方案"],
            ["tau_cross", "0.85", "当前主评价和 E 组 split 采用的阈值"],
            ["Naive union", "image>=0.8 OR text>=0.6", "用于证明 Stage 4 不是单模态结果简单拼接"],
        ],
    )


def add_experiments(doc: Document, data: dict[str, Any]) -> None:
    add_heading(doc, "3. 实验设计与当前数据", 1)
    add_experiment_1(doc, data)
    add_experiment_2(doc, data)
    add_experiment_3(doc, data)
    add_experiment_4(doc, data)
    add_experiment_5(doc, data)
    add_experiment_6(doc, data)


def add_experiment_1(doc: Document, data: dict[str, Any]) -> None:
    labels = data["audit"]["current_key_numbers"]["label_distribution"]
    checks = data["audit"]["hard_consistency_checks"]
    add_heading(doc, "实验 1：CC3M 图文对数据池与人工标注集构建", 2)
    add_experiment_intro(
        doc,
        purpose="构建真实数据上的图文对去重 benchmark，替代原来纯合成重复数据评估，使 Stage 4 的 P/R/F1 来自真实 CC3M 场景。",
        design="先从 CC3M 下载并清洗 200K image-caption pairs，形成可追溯 manifest；再用 image/text/joint 多路相似度挖掘 500K candidate pair-pairs；最后从 high-joint hard candidates 中抽样 1,000 条进行人工标注，保证标注集中有足够重复与近重复样本。",
        params="数据源为 CC3M；当前 pool 为 200K pairs；候选挖掘规模为 500K pair-pairs；人工标签包括 duplicate、near-duplicate、not-duplicate；评价时把 duplicate 与 near-duplicate 合并为 positive。",
        resource="数据准备和候选挖掘主要在 Windows RTX 3090 机器完成；人工标注通过本项目标注前端完成；最终 annotation、metrics 与审计文件同步回 Mac 作为 source-of-truth。",
    )
    add_table(
        doc,
        ["项目", "当前数据", "说明"],
        [
            ["CC3M pool", fmt_int(checks["manifest_rows"]), "已准备 image-caption pairs"],
            ["Candidate pair-pairs", fmt_int(checks["candidate_rows"]), "由 image/text/joint 相似度挖掘"],
            ["High-joint candidates", fmt_int(checks["high_joint_rows"]), "用于人工标注采样"],
            ["已标注 pair-pairs", fmt_int(checks["annotation_rows"]), "当前 Stage 4 GT"],
            ["duplicate", fmt_int(labels["duplicate"]), "严格重复"],
            ["near-duplicate", fmt_int(labels["near_duplicate"]), "语义或视觉近重复"],
            ["not-duplicate", fmt_int(labels["not_duplicate"]), "负例"],
            ["positive total", fmt_int(labels["positive_total"]), "duplicate + near-duplicate"],
        ],
    )


def add_experiment_2(doc: Document, data: dict[str, Any]) -> None:
    best = data["stage4_eval"]["best_by_score"]
    err = data["error_analysis"]
    add_heading(doc, "实验 2：Stage 4 主评价", 2)
    add_experiment_intro(
        doc,
        purpose="回答审稿人最关心的问题：新增的跨模态 Stage 4 是否真的不同于三个单模态模块拼接，尤其是否优于 naive multimodal union。",
        design="在 1,000 条人工标注 hard candidates 上统一计算四类方法：image-only、text-only、naive union 和 Stage 4 joint。image-only 与 text-only 表示单模态上限/边界，naive union 表示直接拼接两个单模态删除集合，Stage 4 joint 表示图文对联合表示判重。四类方法共享同一标注集和同一正负例定义，避免评价口径不一致。",
        params="正例定义为 duplicate + near-duplicate；image-only 当前阈值为 0.8；text-only 当前阈值为 0.6；naive union 为 image>=0.8 OR text>=0.6；Stage 4 当前采用 concat joint embedding，tau_cross = 0.85。",
        resource="评价脚本在 Mac 端基于已同步 annotation、candidate scores 和 metrics source-of-truth 运行；论文中引用的每个数字都应指向对应 metrics 文件和 experiment ledger。",
    )
    add_table(
        doc,
        ["方法", "阈值", "Precision", "Recall", "F1", "TP", "FP", "TN", "FN"],
        [
            eval_row("Image-only", best["image"]),
            eval_row("Text-only", best["text"]),
            eval_row("Naive union", best["naive_union"]),
            eval_row("Stage 4 joint", best["joint"]),
        ],
    )
    add_table(
        doc,
        ["误差项", "当前数据", "解释"],
        [
            ["Stage 4 false positives", fmt_int(err["joint_false_positives"]), "joint 判为重复但人工标签为负例"],
            ["Stage 4 false negatives", fmt_int(err["joint_false_negatives"]), "人工正例中被 joint 阈值漏掉"],
            ["Image correct / joint wrong", fmt_int(err["image_correct_joint_wrong"]), "解释 image-only 当前更强"],
            ["Joint correct / image wrong", fmt_int(err["joint_correct_image_wrong"]), "说明跨模态信号的增量价值"],
            ["Joint FP with identical captions", f"{fmt_int(err['joint_fp_caption_equal'])} ({pct(err['joint_fp_caption_equal_rate'])})", "caption template failure mode"],
        ],
    )


def add_experiment_3(doc: Document, data: dict[str, Any]) -> None:
    add_heading(doc, "实验 3：Stage 4 阈值敏感性", 2)
    add_experiment_intro(
        doc,
        purpose="确定 tau_cross 的可报告工作点，并展示 Stage 4 在 precision、recall、F1 和实际去重率之间的权衡。",
        design="第一步在 1,000 条人工标注集上扫描 joint threshold，观察 P/R/F1 如何变化；第二步把同一批阈值应用到 200K CC3M candidate graph，统计每个阈值对应的 connected components、drops 和 dedup rate。这样可以同时回答“评价集上是否准”和“真实训练池中会删掉多少数据”。",
        params="扫描 tau_cross = 0.60, 0.65, 0.70, 0.75, 0.80, 0.85, 0.90, 0.95；当前选定 0.85，是为了在保留较高 precision 的同时保持可观 recall 和适中的训练数据缩减。",
        resource="评价指标来自 1,000 条标注集；去重率来自 200K candidate graph；阈值扫描结果必须保存完整 CSV，避免只报告 best number。",
    )
    add_table(
        doc,
        ["tau_cross", "Precision", "Recall", "F1", "200K 去重率", "选择说明"],
        joint_threshold_rows(data),
    )


def add_experiment_4(doc: Document, data: dict[str, Any]) -> None:
    splits = data["split_metrics"]["best_known_split_sizes"]
    add_heading(doc, "实验 4：A/B/C/D/E 训练数据构建", 2)
    add_experiment_intro(
        doc,
        purpose="构建下游 LLaVA 对照实验所需的五组训练数据，隔离 raw、单模态去重、naive union 与 Stage 4 对训练效果和训练成本的影响。",
        design="在同一 200K 图文池上，根据不同去重策略生成 A/B/C/D/E 五组 manifest。A 组不去重，用作 raw baseline；B/C 分别只使用图像或文本信号；D 组使用两个单模态删除集合的并集，模拟“多模态就是单模态拼接”的方案；E 组使用 Stage 4 joint threshold，是本轮方法组。后续所有组使用相同 LLaVA LoRA 训练脚本、相同训练步数和相同采样规模。",
        params="A 为 raw；B 使用 image>=0.8；C 使用 text>=0.6；D 使用 image>=0.8 OR text>=0.6；E 使用 joint>=0.85；每组最终训练采样 25K records，目标训练 2,000 steps。",
        resource="split 生成和 manifest 写入在 source-of-truth 中记录；训练在 Windows RTX 3090 上执行；Windows 产物必须同步回 Mac 后才作为稳定论文数字。",
    )
    rows = []
    for split in ["A", "B", "C", "D", "E"]:
        row = splits[split]
        rows.append(
            [
                split,
                row["name"],
                fmt_int(row["raw_pairs"]),
                fmt_int(row["kept_pairs"]),
                fmt_int(row["dropped_pairs"]),
                pct(row["dedup_rate"]),
                row["threshold"],
            ]
        )
    add_table(doc, ["配置", "名称", "原始样本", "保留", "删除", "去重率", "规则"], rows)


def add_experiment_5(doc: Document, data: dict[str, Any]) -> None:
    add_heading(doc, "实验 5：LLaVA 下游训练与 VQAv2 评测", 2)
    add_experiment_intro(
        doc,
        purpose="验证去重后的训练数据是否能够在减少重复样本的同时保持 MLLM 下游能力，重点比较 Stage 4 与 raw、单模态去重和 naive union。",
        design="对 A/B/C/D/E 五组数据分别 fine-tune LLaVA-1.5-7B LoRA。训练阶段记录 loaded records、steps、final loss、wall-clock time 和 GPU peak memory；评测阶段在 VQAv2 上比较 accuracy。训练 loss 只用于观察收敛和排查异常，最终方法有效性应以下游 VQAv2 指标和数据缩减率共同判断。",
        params="每组训练数据采样 25K records，目标 2,000 steps；batch size = 1，gradient accumulation = 8；使用 4-bit LoRA 以适配单张 RTX 3090 24GB；A/B/C/D/E 保持相同训练超参。",
        resource="训练与评测在 Windows RTX 3090 24GB 上执行。当前 A/B/C/D 已完成训练；E 训练结果与 VQAv2 评测结果在表中留空，等待后续实验回填。",
    )
    llava = data["llava"]
    rows = []
    descriptions = {"A": "raw", "B": "image-only", "C": "text-only", "D": "naive union", "E": "Stage 4 joint"}
    for split in ["A", "B", "C", "D", "E"]:
        metric = llava.get(split)
        if metric:
            rows.append(
                [
                    split,
                    descriptions[split],
                    fmt_int(metric.get("num_loaded_records")),
                    metric.get("steps"),
                    f"{metric.get('final_loss'):.4f}",
                    f"{metric.get('runtime_seconds') / 3600:.2f}",
                    f"{metric.get('gpu_peak_memory_bytes') / 1024**3:.3f}",
                ]
            )
        else:
            rows.append([split, descriptions[split], "", "", "", "", ""])
    add_table(doc, ["配置", "名称", "样本数", "训练步数", "最终 loss", "Wall-clock h", "Peak GB"], rows)
    add_table(
        doc,
        ["配置", "训练时间", "相对 A", "平均每 step", "说明"],
        training_time_rows(llava),
    )
    add_paragraph(
        doc,
        "需要注意的是，表中的 final loss 是最后一个 optimizer step 对应 mini-batch 的训练交叉熵，"
        "不是整组数据上的平均 loss，也不是 VQAv2 下游性能。不同 split 的样本排序、caption 长度、"
        "图文难度和最后一个 batch 的组成不同，因此单个 final loss 会有较大波动。比如 B 组最后一步 "
        "loss 较低，并不能直接解释为 B 组下游效果最好。为了更稳妥地观察训练收敛趋势，下面额外给出"
        "前 50 个 step 平均 loss、后 50 个 step 平均 loss，以及二者的比值作为归一化后的相对下降指标。"
    )
    add_table(
        doc,
        ["配置", "前 50 步平均 loss", "后 50 步平均 loss", "归一化 loss 比值", "解释"],
        normalized_loss_rows(llava),
    )
    add_table(
        doc,
        ["配置", "VQAv2 Acc", "评测样本数", "评测设置", "备注"],
        [[split, "", "", "", ""] for split in ["A", "B", "C", "D", "E"]],
    )


def add_experiment_6(doc: Document, data: dict[str, Any]) -> None:
    cc3m = data["cc3m_prepare"]
    mining = data["candidate_mining"]
    split_threshold = data["split_threshold"]
    manifests = data["split_metrics"]
    eval_metrics = data["stage4_eval"]
    add_heading(doc, "实验 6：系统效率与开销", 2)
    add_experiment_intro(
        doc,
        purpose="说明新增 Stage 4 带来的系统开销是否可接受，并区分一次性数据准备成本与算法本身的候选挖掘/图计算成本。",
        design="记录 CC3M 准备、OpenCLIP 候选挖掘、阈值到 split 的 graph computation、A/B/C/D/E manifest 写入和 1K 标注集评价耗时。论文写作时应把数据下载/图片保存和 Stage 4 算法开销分开表述，避免把网络和 I/O 成本误写成算法复杂度。",
        params="候选挖掘使用 200K pairs、500K candidate edges；split 阈值采用当前 A/B/C/D/E 设定；throughput 分别按 pairs/s、edges/s 或 rows/s 计算。",
        resource="候选挖掘在 Windows RTX 3090 上完成；source-of-truth consolidation 在 Mac 上完成；效率表只把 Windows 3090 结果作为主要硬件报告口径。",
    )
    add_table(
        doc,
        ["阶段", "规模", "时间", "Throughput", "说明"],
        [
            ["CC3M 200K prepare", f"{fmt_int(cc3m['saved_pairs'])} pairs", duration(cc3m["elapsed_seconds"]), f"{cc3m['saved_pairs'] / cc3m['elapsed_seconds']:.2f} pairs/s", "数据准备，不计为 Stage 4 算法开销"],
            ["OpenCLIP candidate mining", f"{fmt_int(mining['num_pairs'])} pairs / {fmt_int(mining['num_candidates'])} edges", duration(mining["elapsed_seconds"]), f"{mining['num_pairs'] / mining['elapsed_seconds']:.2f} pairs/s", "Windows RTX 3090"],
            ["Threshold-to-split graph", f"{fmt_int(split_threshold['num_pairs'])} pairs / {fmt_int(split_threshold['num_candidates'])} edges", duration(split_threshold["elapsed_seconds"]), f"{split_threshold['num_candidates'] / split_threshold['elapsed_seconds']:.2f} edges/s", "source-of-truth consolidation"],
            ["A/B/C/D/E manifest writing", f"{fmt_int(manifests['num_pairs'])} pairs", duration(manifests["elapsed_seconds"]), f"{manifests['num_pairs'] / manifests['elapsed_seconds']:.2f} pairs/s", "生成训练 manifest"],
            ["1K labeled evaluation", f"{fmt_int(eval_metrics['num_labeled_rows'])} rows", duration(eval_metrics["elapsed_seconds"]), f"{eval_metrics['num_labeled_rows'] / eval_metrics['elapsed_seconds']:.2f} rows/s", "指标计算"],
        ],
    )


def add_source_rule(doc: Document) -> None:
    doc.add_section(WD_SECTION.CONTINUOUS)
    add_heading(doc, "4. 结果写入规则", 1)
    add_paragraph(
        doc,
        "所有表格数字必须从 source-of-truth 文件写入，包括 experiment_ledger.csv、metrics.json、"
        "per-threshold CSV、训练日志和 Windows 同步目录。E 组训练结果和 VQAv2 评测结果产生后，"
        "只需要回填本文实验 5 的两张表，并重新生成 Word。"
    )


def add_experiment_intro(doc: Document, purpose: str, design: str, params: str, resource: str) -> None:
    for label, text in [
        ("实验目的", purpose),
        ("实验设计", design),
        ("参数设置", params),
        ("资源环境", resource),
    ]:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.18)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(2)
        add_inline_label(para, label, text)


def add_inline_label(para, label: str, text: str) -> None:
    label_run = para.add_run(f"{label}：")
    label_run.bold = True
    set_run_font(label_run)
    text_run = para.add_run(text + " ")
    set_run_font(text_run)


def joint_threshold_rows(data: dict[str, Any]) -> list[list[Any]]:
    dedup = {row["threshold"]: row for row in data["dedup_rows"] if row["score"] == "joint"}
    wanted = {"0.6", "0.65", "0.7", "0.75", "0.8", "0.85", "0.9", "0.95"}
    rows = []
    for row in data["threshold_rows"]:
        if row["score"] != "joint" or row["threshold"] not in wanted:
            continue
        d = dedup[row["threshold"]]
        rows.append(
            [
                row["threshold"],
                pct(float(row["precision"])),
                pct(float(row["recall"])),
                f"{float(row['f1']):.4f}",
                pct(float(d["dedup_rate"])),
                "当前选定" if row["threshold"] == "0.85" else "",
            ]
        )
    return rows


def eval_row(label: str, row: dict[str, Any]) -> list[Any]:
    return [
        label,
        row["threshold"],
        pct(row["precision"]),
        pct(row["recall"]),
        f"{row['f1']:.4f}",
        row["tp"],
        row["fp"],
        row["tn"],
        row["fn"],
    ]


def normalized_loss_rows(llava: dict[str, dict[str, Any]]) -> list[list[Any]]:
    descriptions = {
        "A": "raw",
        "B": "image-only",
        "C": "text-only",
        "D": "naive union",
        "E": "Stage 4 joint",
    }
    rows: list[list[Any]] = []
    for split in ["A", "B", "C", "D", "E"]:
        metric = llava.get(split)
        if not metric or not metric.get("losses"):
            rows.append([f"{split} ({descriptions[split]})", "", "", "", ""])
            continue
        losses = [float(value) for value in metric["losses"]]
        first = statistics.mean(losses[:50])
        last = statistics.mean(losses[-50:])
        ratio = last / first if first else 0.0
        rows.append(
            [
                f"{split} ({descriptions[split]})",
                f"{first:.4f}",
                f"{last:.4f}",
                f"{ratio:.4f}",
                "比值越低表示训练 loss 相对下降越多；仍不能替代 VQAv2 评测。",
            ]
        )
    return rows


def training_time_rows(llava: dict[str, dict[str, Any]]) -> list[list[Any]]:
    descriptions = {
        "A": "raw",
        "B": "image-only",
        "C": "text-only",
        "D": "naive union",
        "E": "Stage 4 joint",
    }
    baseline = llava.get("A", {}).get("runtime_seconds")
    rows: list[list[Any]] = []
    for split in ["A", "B", "C", "D", "E"]:
        metric = llava.get(split)
        if not metric:
            rows.append([f"{split} ({descriptions[split]})", "", "", "", ""])
            continue
        runtime = float(metric["runtime_seconds"])
        steps = float(metric.get("steps") or 0)
        relative = runtime / float(baseline) if baseline else 0.0
        rows.append(
            [
                f"{split} ({descriptions[split]})",
                f"{runtime / 3600:.2f} h ({runtime / 60:.1f} min)",
                f"{relative:.3f}x",
                f"{runtime / steps:.2f} s/step" if steps else "",
                "25K records / 2000 steps / RTX 3090",
            ]
        )
    return rows


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_heading(text, level=level)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(5)


def add_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.18
    run = para.add_run(text)
    set_run_font(run)


def add_note_box(doc: Document, title: str, bullets: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_border(cell, "CBD5E1")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    set_run_font(r)
    for item in bullets:
        p = cell.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.15)
        r = p.add_run("• " + item)
        set_run_font(r)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[Any]], compact: bool = False) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], str(header), bold=True, color="FFFFFF")
        set_cell_shading(hdr[idx], "334155")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            if compact:
                cells[idx].width = Cm(3 if idx == 0 else 13)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell, "CBD5E1")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    r.font.size = Pt(8.5 if not compact else 9)
                    set_run_font(r)
            if row_idx > 0 and row_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
    doc.add_paragraph()


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)


def set_run_font(run) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def fmt_int(value: Any) -> str:
    if value in (None, ""):
        return ""
    return f"{int(value):,}"


def pct(value: float) -> str:
    return f"{float(value) * 100:.2f}%"


def duration(seconds: Any) -> str:
    seconds = float(seconds)
    if seconds < 60:
        return f"{seconds:.2f} s"
    if seconds < 3600:
        return f"{seconds / 60:.2f} min"
    return f"{seconds / 3600:.2f} h"


if __name__ == "__main__":
    raise SystemExit(main())
